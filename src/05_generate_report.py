"""
05_generate_report.py - Professional medical AI report generation.

Required top-level structure:
1. (一) 實驗結果
2. (二) 系統比較
3. (三) 結論
"""
from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from config import FIGURE_DIR, REPORT_DIR, RESULTS_FILE

logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")
LOGGER = logging.getLogger(__name__)

ANALYSIS_SUMMARY_FILE = RESULTS_FILE.parent / "analysis_summary.json"
REPORT_DOCX = REPORT_DIR / "報告.docx"
REPORT_MD = REPORT_DIR / "報告.md"
REPORT_PDF = REPORT_DIR / "報告.pdf"

REPORT_TITLE = "Hypothyroid ML: Multi-Classifier Diagnostic System"
REPORT_SUBTITLE = "Medical AI Benchmarking Report"
SECTION_EXPERIMENT = "(一) 實驗結果"
SECTION_COMPARISON = "(二) 系統比較"
SECTION_CONCLUSION = "(三) 結論"


def load_json(path: Path) -> Any:
    if not path.exists():
        raise FileNotFoundError(f"Missing required file: {path}")
    with path.open("r", encoding="utf-8") as file:
        return json.load(file)


def percent(value: float | None) -> str:
    if value is None:
        return "-"
    return f"{value * 100:.2f}%"


def setup_document_style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = document.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"].font.size = Pt(10.5)
    for level in ["Heading 1", "Heading 2"]:
        styles[level].font.name = "Microsoft JhengHei"
        styles[level].font.bold = True
        styles[level].font.color.rgb = RGBColor(27, 77, 89)


def add_title(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(REPORT_TITLE)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(27, 77, 89)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(REPORT_SUBTITLE)
    run2.font.size = Pt(13)
    run2.font.color.rgb = RGBColor(69, 90, 100)

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("資料來源：hypothyroid_cjlin2025_training.arff / hypothyroid_cjlin2025_test.arff")


def add_abstract(document: Document, best_model: dict[str, Any], dataset: dict[str, Any]) -> None:
    document.add_heading("摘要", level=1)
    document.add_paragraph(
        "本研究以官方測試集為唯一最終評估基準，建立傳統機器學習與 TensorFlow DenseNN "
        "分類系統，用於甲狀腺低下症結構化資料分類。報告重點不僅是排序模型，"
        "而是判斷哪些指標真正反映少數疾病類別的臨床風險。"
    )
    document.add_paragraph(
        f"資料集包含訓練樣本 {dataset['train_samples']} 筆、測試樣本 {dataset['test_samples']} 筆、"
        f"特徵 {dataset['feature_count']} 個。多類別最大/最小類別不平衡比為 "
        f"{dataset['multiclass_imbalance_ratio']:.0f}:1；若轉為臨床篩檢視角，"
        f"negative 對所有 hypothyroid 類別合計約為 {dataset['screening_imbalance_ratio']:.2f}:1。"
    )
    document.add_paragraph(
        f"最佳模型為 {best_model['system']}（{best_model['classifier']}），"
        f"Weighted F1={percent(best_model.get('f1_weighted'))}、"
        f"Balanced Accuracy={percent(best_model.get('balanced_accuracy'))}、"
        f"ROC-AUC={percent(best_model.get('roc_auc_ovr_weighted'))}。"
    )


def add_table(document: Document, results: list[dict[str, Any]]) -> None:
    headers = [
        "系統", "分類器", "Accuracy", "Weighted F1", "Macro F1",
        "Recall", "Balanced Acc", "ROC-AUC", "PR-AUC"
    ]
    table = document.add_table(rows=1, cols=len(headers), style="Table Grid")
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    for row in sorted(results, key=lambda item: item.get("f1_weighted", 0), reverse=True):
        vals = [
            row["system"],
            row["classifier"],
            percent(row.get("accuracy")),
            percent(row.get("f1_weighted")),
            percent(row.get("f1_macro")),
            percent(row.get("recall_macro", row.get("recall"))),
            percent(row.get("balanced_accuracy")),
            percent(row.get("roc_auc_ovr_weighted")),
            percent(row.get("pr_auc_weighted")),
        ]
        cells = table.add_row().cells
        for i, value in enumerate(vals):
            cells[i].text = value


def add_figure(document: Document, name: str, caption: str, width: float = 6.4) -> None:
    path = FIGURE_DIR / name
    if not path.exists():
        LOGGER.warning("Missing figure: %s", path)
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))

    cp = document.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(69, 90, 100)


def add_insight(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(27, 77, 89)


def build_comparison_text(best: dict[str, Any]) -> str:
    return (
        "樹系模型在此資料上的優勢並非單純來自模型容量，而是來自其對表格型臨床資料的歸納偏誤。"
        "TSH、FTI、TT4 等檢驗值具有非線性閾值與交互關係，Gradient Boosting 能逐步修正殘差，"
        "Random Forest 能以多棵樹平均降低變異，因此比單一線性邊界或條件獨立假設更適合此任務。\n\n"
        "DenseNN 的整體準確率看似接近，但 Macro F1 與 Balanced Accuracy 較弱，代表少數類別辨識不足。"
        "在僅約三千筆訓練樣本、特徵數有限且類別極不平衡的情境下，DenseNN 需要更多資料、"
        "更強的校準與更細緻的重抽樣策略，否則容易學到以多數類別為中心的決策表面。\n\n"
        "Weighted 指標會依類別支持度加權，negative 類別占比過高時會放大多數類別的成功；"
        "Macro 指標則讓每個類別擁有相同權重，因此更能暴露 compensated、primary、secondary "
        "hypothyroid 等少數疾病類別的弱點。醫療篩檢情境尤其不能只看 Accuracy，"
        "因為 False Negative 代表可能漏診，臨床風險通常高於可再檢驗的 False Positive。\n\n"
        f"依本次結果，建議優先採用 {best['system']}（{best['classifier']}）作為基準部署版本，"
        "並在部署前進一步完成外部驗證、機率校準、閾值敏感度分析與臨床工作流程測試。"
    )


def build_sensitivity_tradeoff_text(results: list[dict[str, Any]], best: dict[str, Any]) -> str:
    sensitivity_leader = max(results, key=lambda row: row.get("balanced_accuracy", 0.0))
    if sensitivity_leader["system"] == best["system"]:
        return (
            f"{best['system']} 同時是 Weighted F1 與 Balanced Accuracy 的最佳模型，"
            "表示整體準確性與類別平衡召回之間沒有明顯衝突。"
        )
    return (
        f"值得注意的是，{best['system']} 是 Weighted F1 最佳模型，"
        f"但 {sensitivity_leader['system']} 的 Balanced Accuracy 較高"
        f"（{percent(sensitivity_leader.get('balanced_accuracy'))}）。"
        "這代表若系統目標是一般整體分類品質，Gradient Boosting 較穩健；"
        "若臨床場景優先要求降低漏診，則應進一步比較 AdaBoost 的 false positive 成本、"
        "少數類別召回與閾值調整後的可接受性。"
    )


def build_conclusion_text(best: dict[str, Any], feature_summary: dict[str, Any]) -> str:
    top_features = feature_summary.get("top_features", [])[:5]
    feature_text = "、".join([f"{item['feature']}({item['importance']:.3f})" for item in top_features]) if top_features else "無"
    return (
        f"最佳系統設定建議：{best['system']} / {best['classifier']}，參數：{best.get('config', '-') }。\n"
        f"核心指標：Accuracy {percent(best.get('accuracy'))}、Weighted F1 {percent(best.get('f1_weighted'))}、"
        f"Recall {percent(best.get('recall_macro', best.get('recall')))}、Balanced Accuracy {percent(best.get('balanced_accuracy'))}。\n\n"
        "綜合 PR-AUC、類別召回熱圖、複雜度-效能關係、泛化曲線與臨床錯誤分布分析，"
        "模型在測試集展現高整體辨識率，但少數類別仍是臨床導入前必須持續監控的主要風險。\n"
        f"重要特徵（樹模型）前五名：{feature_text}。\n\n"
        "臨床應用上，建議將本系統定位為 AI 輔助判讀工具，而非獨立診斷決策。"
        "後續應加入 SHAP 個案層級解釋、少數類別閾值最佳化、外部資料驗證與資料漂移監測，"
        "以強化可解釋性、跨場域可靠度與臨床安全性。"
    )


def write_markdown_report(results: list[dict[str, Any]], summary: dict[str, Any]) -> None:
    best = summary["best_model"]
    dataset = summary["dataset"]
    lines: list[str] = [
        f"# {REPORT_TITLE}",
        "",
        "## 摘要",
        f"- 訓練集: {dataset['train_samples']}，測試集: {dataset['test_samples']}，特徵: {dataset['feature_count']}",
        f"- 多類別最大/最小不平衡比: {dataset['multiclass_imbalance_ratio']:.0f}:1；篩檢視角 negative:positive = {dataset['screening_imbalance_ratio']:.2f}:1",
        f"- 最佳模型: {best['system']} ({best['classifier']})",
        f"- Weighted F1: {percent(best.get('f1_weighted'))}；Balanced Accuracy: {percent(best.get('balanced_accuracy'))}",
        "",
        f"## {SECTION_EXPERIMENT}",
        "| System | Classifier | Acc | F1w | MacroF1 | Recall | BAcc | ROC-AUC | PR-AUC |",
        "|---|---|---:|---:|---:|---:|---:|---:|---:|",
    ]

    for row in sorted(results, key=lambda item: item.get("f1_weighted", 0), reverse=True):
        lines.append(
            f"| {row['system']} | {row['classifier']} | {percent(row.get('accuracy'))} | {percent(row.get('f1_weighted'))} | "
            f"{percent(row.get('f1_macro'))} | {percent(row.get('recall_macro', row.get('recall')))} | "
            f"{percent(row.get('balanced_accuracy'))} | {percent(row.get('roc_auc_ovr_weighted'))} | {percent(row.get('pr_auc_weighted'))} |"
        )

    lines.extend([
        "",
        f"## {SECTION_COMPARISON}",
        build_sensitivity_tradeoff_text(results, best),
        "",
        build_comparison_text(best),
        "",
        f"## {SECTION_CONCLUSION}",
        build_conclusion_text(best, summary.get("feature_importance", {})),
    ])

    REPORT_MD.write_text("\n".join(lines), encoding="utf-8")


def generate_report() -> None:
    results = load_json(RESULTS_FILE)
    summary = load_json(ANALYSIS_SUMMARY_FILE)

    best = summary["best_model"]
    dataset = summary["dataset"]

    doc = Document()
    setup_document_style(doc)
    add_title(doc)
    add_abstract(doc, best, dataset)

    doc.add_heading(SECTION_EXPERIMENT, level=1)
    doc.add_paragraph(
        "所有最終指標均以官方 test set 計算。資料分布顯示 secondary hypothyroid 極罕見，"
        "因此單純 Accuracy 會高估臨床可用性；本報告同步呈現 Weighted、Macro、Balanced "
        "與 PR-AUC 指標，以區分多數類別表現與少數疾病類別風險。"
    )
    add_figure(
        doc,
        "dataset_imbalance.png",
        f"圖1. 類別不平衡不是背景資訊，而是本任務的主要建模條件；多類別 "
        f"{dataset['multiclass_imbalance_ratio']:.0f}:1 與篩檢視角 "
        f"{dataset['screening_imbalance_ratio']:.2f}:1 應分開解讀。",
    )
    add_insight(doc, "重點：Weighted F1 高不代表每個疾病亞型都被穩定辨識，必須搭配 Macro F1 與 Balanced Accuracy。")
    add_table(doc, results)
    add_insight(doc, build_sensitivity_tradeoff_text(results, best))
    add_figure(doc, "metric_fingerprint.png", "圖2. 指標指紋圖揭示 Weighted 與 Macro 指標落差；落差越大，越可能代表少數類別被平均值掩蓋。")
    add_figure(doc, "class_recall_heatmap.png", "圖3. 類別召回熱圖將臨床漏診風險直接映射到各模型與各疾病亞型。")
    add_figure(doc, "pr_multi_model.png", "圖4. Precision-Recall 曲線比 ROC 更適合不平衡任務，可觀察高召回區間是否仍維持可接受精確率。")

    doc.add_heading(SECTION_COMPARISON, level=1)
    doc.add_paragraph(build_comparison_text(best))
    add_figure(doc, "complexity_performance.png", "圖5. 模型複雜度與 Balanced Accuracy 的關係顯示，最佳臨床指標不必然來自最大模型。")
    add_figure(doc, "feature_importance.png", "圖6. 樹模型重要特徵集中於 TSH、FTI、TT4，符合甲狀腺功能判讀的臨床直覺。")
    add_figure(doc, "learning_curve_best_model.png", "圖7. 學習曲線用於判讀泛化差距；若驗證曲線趨於平台，新增資料或少數類別增強比單純加深模型更重要。")
    add_figure(doc, "error_distribution.png", "圖8. 臨床錯誤面板同時呈現 False Positive、False Negative 與正規化混淆流，優先暴露漏診風險。")

    doc.add_heading(SECTION_CONCLUSION, level=1)
    doc.add_paragraph(build_conclusion_text(best, summary.get("feature_importance", {})))

    doc.save(REPORT_DOCX)
    write_markdown_report(results, summary)

    # Optional PDF export if docx2pdf is available on user machine.
    try:
        from docx2pdf import convert  # type: ignore

        convert(str(REPORT_DOCX), str(REPORT_PDF))
        LOGGER.info("PDF exported: %s", REPORT_PDF)
    except Exception:
        LOGGER.info("PDF export skipped (install docx2pdf + Word if PDF is required).")

    LOGGER.info("Report generated: %s", REPORT_DOCX)
    LOGGER.info("Markdown report generated: %s", REPORT_MD)


if __name__ == "__main__":
    generate_report()
